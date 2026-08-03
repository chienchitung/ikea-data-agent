import os
import sys
from pathlib import Path

os.environ.setdefault("GOOGLE_API_KEY", "dummy")

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

import io
from docx import Document as DocxDocument

from agents.meeting import build_docx


SAMPLE_MEETING_DATA = {
    "meeting_title": "Weekly Data & Sharing Discussion",
    "date": "Jan 29, 2025",
    "time": "2pm-3pm",
    "note_taker": "Jackie",
    "attendees": "Jackie, Kelly, Lillian",
    "apologies": "",
    "last_time_actions": [],
    "agenda": [
        {"no": 1, "item": "Review last week's dashboard issue", "owner": "Kelly"},
    ],
    "notes": [
        "Dashboard refresh was delayed due to a schema change in the Request worksheet.",
        "Team agreed to add a validation step before next release.",
    ],
    "actions_review": [
        {"no": 1, "item": "Fix dashboard refresh schedule", "assigned_to": "Kelly", "deadline": "2025-02-05"},
    ],
}


def _table_texts(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def test_build_docx_produces_valid_document():
    docx_bytes = build_docx(SAMPLE_MEETING_DATA)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0

    doc = DocxDocument(io.BytesIO(docx_bytes))
    assert doc.paragraphs[0].text == "Weekly Data & Sharing Discussion"

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Last time actions review" in headings
    assert "Meeting Agenda" in headings
    assert "Meeting Notes" in headings
    assert "Actions review" in headings

    # 3 tables in document order: last-time-actions, agenda, actions-review
    assert len(doc.tables) == 3

    last_time_table = _table_texts(doc.tables[0])
    assert last_time_table[0] == ["No.", "Action items", "Assigned to", "Feedback"]
    # left blank per product decision, but padded to at least 3 rows like the paper template
    assert len(last_time_table) - 1 >= 3

    agenda_table = _table_texts(doc.tables[1])
    assert agenda_table[0] == ["No.", "Agenda items", "Presenter/Owner"]
    assert agenda_table[1][1] == "Review last week's dashboard issue"

    actions_table = _table_texts(doc.tables[2])
    assert actions_table[0] == ["No.", "Action items", "Assigned to", "Deadline"]
    assert actions_table[1][1] == "Fix dashboard refresh schedule"
    assert actions_table[1][2] == "Kelly"


def test_build_docx_handles_empty_meeting_data():
    docx_bytes = build_docx({})
    doc = DocxDocument(io.BytesIO(docx_bytes))
    assert doc.paragraphs[0].text == "Meeting Notes"
    assert len(doc.tables) == 3


if __name__ == "__main__":
    test_build_docx_produces_valid_document()
    test_build_docx_handles_empty_meeting_data()
    print("✅ meeting_docx_smoke passed")
