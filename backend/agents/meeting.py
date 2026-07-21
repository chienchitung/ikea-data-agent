import asyncio
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Awaitable, Callable, Optional

import httpx
from dotenv import load_dotenv

try:
    from pydub import AudioSegment
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False

try:
    from google import genai
    from google.genai import types as genai_types
    from google.genai.errors import ServerError as GenaiServerError
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False

try:
    from opencc import OpenCC
    # Whisper-family models (including Groq's) transcribe Mandarin speech as
    # Simplified characters regardless of the speaker's accent/region, so
    # Traditional-Chinese-speaking teams get output that reads wrong. s2twp
    # converts Simplified -> Traditional with Taiwan phrasing conventions.
    _S2TWP_CONVERTER = OpenCC("s2twp")
    OPENCC_AVAILABLE = True
except ImportError:
    _S2TWP_CONVERTER = None
    OPENCC_AVAILABLE = False

from docx import Document as DocxDocument

load_dotenv()

# Used only for generate_minutes() (text-only) now — transcription is handled
# by Groq's Whisper API below. Not chained through GEMINI_MODEL (the chat
# coordinator's model) so it can be tuned independently: gemini-3.5-flash has
# been observed under heavy load taking 90s+ for a trivial text-only call
# vs. gemini-2.5-flash's ~7s for the same schema request — a known risk we're
# accepting here in exchange for gemini-3.5-flash's output quality. If minutes
# generation starts timing out or feels slow in practice, drop MEETING_MODEL
# back to gemini-2.5-flash via env var rather than reverting this default.
MEETING_MODEL = os.getenv("MEETING_MODEL", "gemini-3.5-flash")
NORMALIZED_AUDIO_FORMAT = "mp3"

# gemini-3.5-flash has been observed returning a transient 503 ("high demand")
# on large-transcript requests (a ~73k-character transcript from a 2.5-hour
# meeting hit this in practice) — without a retry, transcription work that
# already completed is thrown away and the whole meeting is lost. Backoff is
# linear (15s, 30s, 45s) rather than exponential since these are reported as
# short-lived demand spikes, not sustained outages worth minutes-long backoff.
MINUTES_GENERATION_MAX_ATTEMPTS = 4
MINUTES_GENERATION_RETRY_BASE_SECONDS = 15

# Groq runs Whisper on LPU hardware at roughly 200x real-time (a 1-hour
# recording transcribes in ~15-20s), which is why transcription is delegated
# here instead of sending the whole audio file to Gemini as a multimodal
# input — audio ingestion by a general-purpose LLM was the dominant latency
# cost in this pipeline (minutes, not seconds, for a typical meeting).
GROQ_TRANSCRIPTION_MODEL = os.getenv("GROQ_WHISPER_MODEL", "whisper-large-v3-turbo")
GROQ_TRANSCRIPTIONS_URL = "https://api.groq.com/openai/v1/audio/transcriptions"

# Groq caps direct (multipart) file uploads at 25MB on the free tier — and
# even the paid tier's 100MB cap reportedly requires URL-based upload rather
# than posting the file directly, per Groq's own community forum. A 1-hour
# meeting normalized to mp3 is already ~55MB at a typical 128kbps, so this is
# routinely hit rather than an edge case. Chunk size is picked as a safety
# margin under the 25MB limit so this also works unmodified on the free tier.
GROQ_MAX_UPLOAD_BYTES = 20 * 1024 * 1024

BACKEND_DIR = Path(__file__).resolve().parent.parent
from storage_paths import MEETING_STORE_DIR as STORE_DIR, MEETING_AUDIO_DIR as AUDIO_DIR

DEFAULT_MEETING_TITLE = "Meeting Notes"

MINUTES_PROMPT_TEMPLATE = """你是 IKEA Data Team 的會議記錄整理助理。請根據下方的會議逐字稿，整理成結構化的會議記錄 JSON。

輸出規則：
- 只能輸出 JSON，不要輸出任何其他文字或 Markdown。
- "executive_summary"、"agenda"、"notes"、"actions_review" 裡的文字內容，語言必須跟隨逐字稿本身使用的語言；
  逐字稿是英文就用英文寫，逐字稿是中文就用中文寫，中英夾雜就維持中英夾雜，不要自動翻譯成其他語言。
  寫中文時一律使用繁體中文（台灣用語習慣），不要使用簡體字。
- "executive_summary" 是整場會議的總結摘要：用 3~6 句完整敘述先交代會議目的、主要討論脈絡、
  達成的結論與最重要的後續方向，讓沒參加會議的人只讀這段就能掌握全貌。不要用條列式。
- "agenda" 是這場會議實際討論到的議程項目，依討論順序列出；如果逐字稿沒有明確的時長，duration 留空字串。
- "notes" 是這場會議的重點討論內容，用條列式整理，每一條是一個獨立重點、決議或討論細節，愈詳細愈好，
  不要只寫一句空泛的話。
- "actions_review" 是這場會議中新產生或提到的待辦事項/行動項目，包含負責人與期限（如果逐字稿有提到）；
  如果逐字稿沒有明確提到負責人或期限，該欄位留空字串。
- 不要杜撰逐字稿沒有提到的人名、日期或數字。

輸出 JSON schema：
{{
  "executive_summary": "...",
  "agenda": [{{"duration": "", "item": "", "owner": ""}}],
  "notes": ["..."],
  "actions_review": [{{"no": 1, "item": "", "assigned_to": "", "deadline": ""}}]
}}

會議逐字稿：
{transcript}
"""


# ── Structured output schemas ────────────────────────────────
# Using an explicit response_schema (rather than relying on the prompt's JSON
# description alone) makes Gemini's structured output far less likely to drift
# from the expected shape than response_mime_type="application/json" alone.
if GENAI_AVAILABLE:
    _AGENDA_ITEM_SCHEMA = genai_types.Schema(
        type=genai_types.Type.OBJECT,
        properties={
            "duration": genai_types.Schema(type=genai_types.Type.STRING),
            "item": genai_types.Schema(type=genai_types.Type.STRING),
            "owner": genai_types.Schema(type=genai_types.Type.STRING),
        },
        required=["item"],
    )

    _ACTION_ITEM_SCHEMA = genai_types.Schema(
        type=genai_types.Type.OBJECT,
        properties={
            "no": genai_types.Schema(type=genai_types.Type.INTEGER),
            "item": genai_types.Schema(type=genai_types.Type.STRING),
            "assigned_to": genai_types.Schema(type=genai_types.Type.STRING),
            "deadline": genai_types.Schema(type=genai_types.Type.STRING),
        },
        required=["item"],
    )

    MINUTES_RESPONSE_SCHEMA = genai_types.Schema(
        type=genai_types.Type.OBJECT,
        properties={
            "executive_summary": genai_types.Schema(type=genai_types.Type.STRING),
            "agenda": genai_types.Schema(type=genai_types.Type.ARRAY, items=_AGENDA_ITEM_SCHEMA),
            "notes": genai_types.Schema(type=genai_types.Type.ARRAY, items=genai_types.Schema(type=genai_types.Type.STRING)),
            "actions_review": genai_types.Schema(type=genai_types.Type.ARRAY, items=_ACTION_ITEM_SCHEMA),
        },
        required=["executive_summary", "agenda", "notes", "actions_review"],
    )


# ── Audio handling ───────────────────────────────────────────
#
# normalize_audio() and the large-file split in transcribe_audio() both shell
# out to the ffmpeg/ffprobe binaries directly rather than going through pydub's
# AudioSegment. pydub reads a decoded file's *entire* raw PCM into the Python
# process's memory (e.g. a 1-hour recording decodes to several hundred MB of
# PCM) — fine on a dev machine with several GB of RAM, but on Render's free
# tier (512MB, shared with this process's baseline numpy/pandas/faiss/langchain
# footprint) a long meeting recording reliably exceeds the memory limit and
# gets the whole service OOM-killed and restarted mid-request. Invoking ffmpeg
# as a subprocess instead keeps decoding inside ffmpeg's own bounded streaming
# buffers — peak added memory in this process stays in the tens of MB range
# regardless of recording length. pydub is kept only as a last-resort fallback
# for the (unexpected, since pydub itself also requires ffmpeg to do anything
# useful) case where the ffmpeg/ffprobe binaries aren't on PATH.

FFMPEG_BIN = shutil.which("ffmpeg")
FFPROBE_BIN = shutil.which("ffprobe")


def _probe_duration_seconds(path: str) -> float:
    """Returns 0.0 if ffprobe is unavailable or the probe fails."""
    if not FFPROBE_BIN:
        return 0.0
    try:
        result = subprocess.run(
            [FFPROBE_BIN, "-v", "error", "-show_entries", "format=duration",
             "-of", "csv=p=0", path],
            check=True, capture_output=True, text=True,
        )
        return float(result.stdout.strip())
    except Exception:
        return 0.0


def _run_ffmpeg_with_progress(cmd: list, total_seconds: float, on_progress: Optional[Callable[[int], None]] = None) -> None:
    """
    Runs an ffmpeg command — which must include "-progress pipe:1 -nostats" —
    and parses its machine-readable progress stream (out_time_us=... lines) to
    report integer percent-complete via on_progress as the transcode/split
    actually runs, instead of the caller only finding out once the whole
    (possibly tens-of-seconds-long, for a long meeting) subprocess call
    returns. `on_progress` must be synchronous and safe to call from this
    thread — this function itself is a normal blocking call.

    stderr is captured to a temp file rather than a pipe: ffmpeg can write
    enough to it that an undrained PIPE fills up and deadlocks the process
    while we're only reading stdout.
    """
    with tempfile.TemporaryFile() as stderr_capture:
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=stderr_capture, text=True, bufsize=1)
        last_percent = -1
        try:
            for line in process.stdout:
                line = line.strip()
                if on_progress and total_seconds > 0 and line.startswith("out_time_us="):
                    raw_value = line.split("=", 1)[1]
                    if raw_value == "N/A":
                        continue
                    elapsed_seconds = max(0, int(raw_value)) / 1_000_000
                    percent = min(99, int(elapsed_seconds / total_seconds * 100))
                    if percent != last_percent:
                        last_percent = percent
                        on_progress(percent)
                elif on_progress and line == "progress=end" and last_percent != 100:
                    last_percent = 100
                    on_progress(100)
        finally:
            process.wait()

        if process.returncode != 0:
            stderr_capture.seek(0)
            stderr_text = stderr_capture.read().decode("utf-8", errors="replace")
            raise subprocess.CalledProcessError(process.returncode, cmd, stderr=stderr_text)


def normalize_audio(src_path: str, on_progress: Optional[Callable[[int], None]] = None) -> str:
    """
    Convert an uploaded/recorded audio file to a Gemini-supported format (mp3).
    Browser MediaRecorder produces audio/webm;codecs=opus, which is not in Gemini's
    documented supported audio MIME list (wav/mp3/aiff/aac/ogg/flac). Falls back to
    the original file if neither ffmpeg nor pydub is available so the pipeline
    still attempts transcription with whatever format was uploaded.

    `on_progress(percent)` is called as the ffmpeg transcode runs (see
    _run_ffmpeg_with_progress) when using the ffmpeg path; the pydub fallback
    has no equivalent hook since pydub only returns once fully done.
    """
    dst_path = f"{os.path.splitext(src_path)[0]}_normalized.{NORMALIZED_AUDIO_FORMAT}"

    if FFMPEG_BIN:
        try:
            total_seconds = _probe_duration_seconds(src_path)
            _run_ffmpeg_with_progress(
                [
                    FFMPEG_BIN, "-y", "-i", src_path, "-vn", "-acodec", "libmp3lame", "-b:a", "128k",
                    "-progress", "pipe:1", "-nostats",
                    dst_path,
                ],
                total_seconds, on_progress,
            )
            return dst_path
        except Exception as e:
            print(f"⚠️ ffmpeg audio normalization failed ({e}); falling back to pydub.")

    if not PYDUB_AVAILABLE:
        print("⚠️ pydub not installed; sending original audio file to Gemini as-is.")
        return src_path
    try:
        audio = AudioSegment.from_file(src_path)
        audio.export(dst_path, format=NORMALIZED_AUDIO_FORMAT)
        return dst_path
    except Exception as e:
        print(f"⚠️ Audio normalization failed ({e}); sending original file to Gemini as-is.")
        return src_path


# ── Structured minutes generation ────────────────────────────

def _parse_json_object(text: str) -> dict:
    try:
        return json.loads(text)
    except Exception:
        pass
    match = re.search(r"\{.*\}", str(text or ""), re.DOTALL)
    if not match:
        return {}
    try:
        return json.loads(match.group(0))
    except Exception:
        return {}


def _format_timestamp(total_seconds) -> str:
    total = max(0, int(total_seconds or 0))
    minutes, seconds = divmod(total, 60)
    return f"[{minutes:02d}:{seconds:02d}]"


def _normalize_segments(items) -> list:
    if not isinstance(items, list):
        return []
    normalized = []
    for item in items:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text", "") or "").strip()
        if not text:
            continue
        try:
            start = float(item.get("start", 0) or 0)
        except (TypeError, ValueError):
            start = 0.0
        normalized.append({
            "start": start,
            "speaker": str(item.get("speaker", "") or "").strip(),
            "text": text,
        })
    normalized.sort(key=lambda seg: seg["start"])
    return normalized


def _to_traditional(text) -> str:
    text = str(text or "")
    if not text or not OPENCC_AVAILABLE:
        return text
    return _S2TWP_CONVERTER.convert(text)


def _segments_to_text(segments: list) -> str:
    lines = []
    for seg in segments:
        prefix = _format_timestamp(seg.get("start"))
        speaker = seg.get("speaker")
        if speaker:
            prefix += f" {speaker}:"
        lines.append(f"{prefix} {seg.get('text', '')}".strip())
    return "\n".join(lines)


async def _transcribe_bytes(audio_bytes: bytes, filename: str, groq_api_key: str) -> list:
    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.post(
            GROQ_TRANSCRIPTIONS_URL,
            headers={"Authorization": f"Bearer {groq_api_key}"},
            files={"file": (filename, audio_bytes)},
            data={
                "model": GROQ_TRANSCRIPTION_MODEL,
                "response_format": "verbose_json",
            },
        )
    response.raise_for_status()
    return response.json().get("segments") or []


def _split_audio_ffmpeg(
    path: str, chunk_seconds: int, out_dir: str,
    total_seconds: float = 0.0, on_progress: Optional[Callable[[int], None]] = None,
) -> list:
    """
    Splits `path` into sequential chunk files under `out_dir` using ffmpeg's
    segment muxer with stream copy (-c copy) — i.e. it repackages existing
    encoded frames into separate files rather than decoding+re-encoding, so it
    doesn't pull the audio through Python at all and runs in roughly constant
    memory regardless of file length. Requires `path` to already be the format
    produced by normalize_audio() (mp3), which segments cleanly with copy.

    `on_progress(percent)` is called as the split runs (see
    _run_ffmpeg_with_progress); `total_seconds` should be this file's already-
    known duration so percent can be computed without a second probe.
    """
    pattern = os.path.join(out_dir, f"chunk_%03d.{NORMALIZED_AUDIO_FORMAT}")
    _run_ffmpeg_with_progress(
        [
            FFMPEG_BIN, "-y", "-i", path,
            "-f", "segment", "-segment_time", str(chunk_seconds),
            "-c", "copy", "-reset_timestamps", "1",
            "-progress", "pipe:1", "-nostats",
            pattern,
        ],
        total_seconds, on_progress,
    )
    return sorted(
        os.path.join(out_dir, name)
        for name in os.listdir(out_dir)
        if name.startswith("chunk_")
    )


def _split_audio_pydub(audio_path: str, file_size: int) -> list:
    """
    Fallback for the (unexpected) case where ffmpeg isn't on PATH. Decodes the
    whole file into memory via pydub, same as the pipeline used to do
    unconditionally — kept only as a last resort since pydub itself normally
    also requires ffmpeg to read anything but raw wav.
    Returns [(offset_seconds, chunk_bytes), ...].
    """
    audio = AudioSegment.from_file(audio_path)
    duration_ms = len(audio)
    bytes_per_ms = file_size / duration_ms if duration_ms else 0
    chunk_ms = int(GROQ_MAX_UPLOAD_BYTES / bytes_per_ms) if bytes_per_ms else duration_ms
    chunk_ms = max(chunk_ms, 60_000)  # never split into slivers under a minute

    chunks = []
    for offset_ms in range(0, duration_ms, chunk_ms):
        buffer = io.BytesIO()
        audio[offset_ms:offset_ms + chunk_ms].export(buffer, format=NORMALIZED_AUDIO_FORMAT)
        chunks.append((offset_ms / 1000.0, buffer.getvalue()))
    return chunks


ProgressCallback = Callable[[int, int], Awaitable[None]]


async def transcribe_audio(
    audio_path: str,
    groq_api_key: str,
    on_chunk_done: Optional[ProgressCallback] = None,
    on_split_progress: Optional[Callable[[int], None]] = None,
) -> dict:
    """
    Returns {"segments": [{"start": seconds, "speaker": str, "text": str}, ...], "text": flat_text}.
    Segments power the clickable, timestamped transcript UI; "text" is a flat
    rendering of the same content used as input to generate_minutes().

    "speaker" is always "" here: Whisper transcribes speech but does not perform
    speaker diarization (unlike the previous audio-aware Gemini prompt, which
    could pick up on distinct voices). The transcript UI already renders the
    speaker label conditionally, so this degrades gracefully.

    `on_chunk_done(completed, total)` is awaited after each chunk finishes when
    the file needed to be split (over Groq's upload limit), so callers can
    surface real transcription progress instead of one static "transcribing"
    status for however long a long meeting takes.

    `on_split_progress(percent)` — synchronous, not awaited — is called while
    ffmpeg is splitting an oversized file into chunks, before any chunk is
    done yet. It runs off the event loop (see asyncio.to_thread below), so
    unlike on_chunk_done it can't safely be an `await`-based callback itself;
    callers bridge it back to async code (e.g. via loop.call_soon_threadsafe).
    """
    file_size = os.path.getsize(audio_path)

    if file_size <= GROQ_MAX_UPLOAD_BYTES:
        with open(audio_path, "rb") as f:
            audio_bytes = f.read()
        raw_segments = await _transcribe_bytes(audio_bytes, os.path.basename(audio_path), groq_api_key)
    elif FFMPEG_BIN:
        duration_seconds = _probe_duration_seconds(audio_path)
        bytes_per_second = file_size / duration_seconds if duration_seconds else 0
        chunk_seconds = int(GROQ_MAX_UPLOAD_BYTES / bytes_per_second) if bytes_per_second else int(duration_seconds) or 60
        chunk_seconds = max(chunk_seconds, 60)  # never split into slivers under a minute

        raw_segments = []
        with tempfile.TemporaryDirectory(prefix="meeting_chunks_") as tmp_dir:
            # Splitting a long file (even with -c copy) can take a few seconds;
            # run it off the event loop so it doesn't stall other requests,
            # consistent with how normalize_audio() is dispatched by the caller.
            chunk_paths = await asyncio.to_thread(
                _split_audio_ffmpeg, audio_path, chunk_seconds, tmp_dir, duration_seconds, on_split_progress,
            )
            total_chunks = len(chunk_paths)
            offset_seconds = 0.0
            # Chunks are transcribed one at a time (not concurrently): only one
            # chunk's bytes + in-flight HTTP request are ever held in memory at
            # once, trading some wall-clock time for a bounded memory peak.
            for i, chunk_path in enumerate(chunk_paths, start=1):
                with open(chunk_path, "rb") as f:
                    chunk_bytes = f.read()
                # Probe before deleting: actual segment length can differ slightly
                # from the requested chunk_seconds since -c copy cuts at the
                # nearest frame boundary, and errors compound across chunks.
                actual_chunk_seconds = _probe_duration_seconds(chunk_path) or chunk_seconds
                os.remove(chunk_path)

                chunk_segments = await _transcribe_bytes(chunk_bytes, os.path.basename(chunk_path), groq_api_key)
                for seg in chunk_segments:
                    raw_segments.append({**seg, "start": (seg.get("start") or 0) + offset_seconds})

                offset_seconds += actual_chunk_seconds
                if on_chunk_done:
                    await on_chunk_done(i, total_chunks)
    else:
        if not PYDUB_AVAILABLE:
            raise RuntimeError(
                f"Audio file is {file_size / (1024 * 1024):.1f}MB, over Groq's per-request upload "
                "limit, and neither ffmpeg nor pydub is available to split it into smaller chunks."
            )
        chunks = _split_audio_pydub(audio_path, file_size)
        raw_segments = []
        for i, (offset_seconds, chunk_bytes) in enumerate(chunks, start=1):
            chunk_segments = await _transcribe_bytes(chunk_bytes, f"chunk_{i}.{NORMALIZED_AUDIO_FORMAT}", groq_api_key)
            for seg in chunk_segments:
                raw_segments.append({**seg, "start": (seg.get("start") or 0) + offset_seconds})
            if on_chunk_done:
                await on_chunk_done(i, len(chunks))

    segments = _normalize_segments([
        {"start": seg.get("start"), "speaker": "", "text": _to_traditional(seg.get("text"))}
        for seg in raw_segments
    ])
    return {
        "segments": segments,
        "text": _segments_to_text(segments),
    }


def _normalize_agenda(items) -> list:
    if not isinstance(items, list):
        return []
    normalized = []
    for item in items:
        if not isinstance(item, dict):
            continue
        normalized.append({
            "duration": _to_traditional(item.get("duration", "")).strip(),
            "item": _to_traditional(item.get("item", "")).strip(),
            "owner": _to_traditional(item.get("owner", "")).strip(),
        })
    return normalized


def _normalize_notes(items) -> list:
    if not isinstance(items, list):
        return []
    return [_to_traditional(item).strip() for item in items if str(item or "").strip()]


def _normalize_actions(items) -> list:
    if not isinstance(items, list):
        return []
    normalized = []
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        normalized.append({
            "no": item.get("no", index),
            "item": _to_traditional(item.get("item", "")).strip(),
            "assigned_to": _to_traditional(item.get("assigned_to", "")).strip(),
            "deadline": _to_traditional(item.get("deadline", "")).strip(),
        })
    return normalized


async def generate_minutes(transcript: str, meta: dict, api_key: str) -> dict:
    if not GENAI_AVAILABLE:
        raise RuntimeError("google-genai SDK is not installed; cannot generate meeting minutes.")

    client = genai.Client(api_key=api_key)
    prompt = MINUTES_PROMPT_TEMPLATE.format(transcript=transcript[:60000])

    response = None
    for attempt in range(1, MINUTES_GENERATION_MAX_ATTEMPTS + 1):
        try:
            response = await client.aio.models.generate_content(
                model=MEETING_MODEL,
                contents=[prompt],
                config={
                    "response_mime_type": "application/json",
                    "response_schema": MINUTES_RESPONSE_SCHEMA,
                },
            )
            break
        except GenaiServerError:
            if attempt == MINUTES_GENERATION_MAX_ATTEMPTS:
                raise
            await asyncio.sleep(MINUTES_GENERATION_RETRY_BASE_SECONDS * attempt)

    parsed = _parse_json_object(str(getattr(response, "text", "") or ""))

    return {
        "meeting_title": str(meta.get("meeting_title", "") or "").strip(),
        "date": str(meta.get("date", "") or "").strip(),
        "time": str(meta.get("time", "") or "").strip(),
        "note_taker": str(meta.get("note_taker", "") or "").strip(),
        "attendees": str(meta.get("attendees", "") or "").strip(),
        "apologies": str(meta.get("apologies", "") or "").strip(),
        # Intentionally left blank: the user fills this in by hand from the
        # previous meeting's "Actions review", same as the paper template.
        "last_time_actions": [],
        "executive_summary": str(parsed.get("executive_summary") or "").strip(),
        "agenda": _normalize_agenda(parsed.get("agenda")),
        "notes": _normalize_notes(parsed.get("notes")),
        "actions_review": _normalize_actions(parsed.get("actions_review")),
    }


# ── .docx generation ─────────────────────────────────────────

def _add_header_field(doc: DocxDocument, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    run_label = paragraph.add_run(f"{label}: ")
    run_label.bold = True
    paragraph.add_run(value or "")


def _add_table(doc: DocxDocument, headers: list, rows: list, min_rows: int = 0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_values in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = str(value or "")

    for _ in range(max(0, min_rows - len(rows))):
        table.add_row()


def build_docx(meeting_data: dict) -> bytes:
    doc = DocxDocument()

    title = meeting_data.get("meeting_title") or DEFAULT_MEETING_TITLE
    doc.add_heading(title, level=1)

    _add_header_field(doc, "Date", meeting_data.get("date", ""))
    _add_header_field(doc, "Time", meeting_data.get("time", ""))
    _add_header_field(doc, "Note taker", meeting_data.get("note_taker", ""))
    _add_header_field(doc, "In attendance", meeting_data.get("attendees", ""))
    _add_header_field(doc, "Apologies for absence", meeting_data.get("apologies", ""))

    doc.add_heading("Last time actions review", level=2)
    last_time_actions = meeting_data.get("last_time_actions") or []
    last_time_rows = [
        [item.get("no", idx + 1), item.get("item", ""), item.get("assigned_to", ""), item.get("feedback", "")]
        for idx, item in enumerate(last_time_actions)
    ]
    _add_table(doc, ["No.", "Action items", "Assigned to", "Feedback"], last_time_rows, min_rows=3)

    doc.add_heading("Meeting Agenda", level=2)
    agenda = meeting_data.get("agenda") or []
    agenda_rows = [[item.get("duration", ""), item.get("item", ""), item.get("owner", "")] for item in agenda]
    _add_table(doc, ["Duration", "Agenda items", "Presenter/Owner"], agenda_rows, min_rows=3)

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(str(meeting_data.get("executive_summary") or ""))

    doc.add_heading("Meeting Notes", level=2)
    notes = meeting_data.get("notes") or []
    if notes:
        for note in notes:
            doc.add_paragraph(str(note), style="List Bullet")
    else:
        doc.add_paragraph("")

    doc.add_heading("Actions review", level=2)
    actions = meeting_data.get("actions_review") or []
    actions_rows = [
        [item.get("no", idx + 1), item.get("item", ""), item.get("assigned_to", ""), item.get("deadline", "")]
        for idx, item in enumerate(actions)
    ]
    _add_table(doc, ["No.", "Action items", "Assigned to", "Deadline"], actions_rows, min_rows=4)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Storage (flat-file-per-ID, mirrors conversation_store.py) ───

def _safe_id(value: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", str(value or ""))
    return safe[:80] or "meeting"


def new_meeting_id() -> str:
    return uuid.uuid4().hex


def audio_dir_for(meeting_id: str) -> Path:
    return AUDIO_DIR / _safe_id(meeting_id)


def _record_path(meeting_id: str) -> Path:
    return STORE_DIR / f"{_safe_id(meeting_id)}.json"


def save_meeting_record(record: dict) -> None:
    STORE_DIR.mkdir(parents=True, exist_ok=True)
    path = _record_path(record["meeting_id"])
    try:
        with path.open("w", encoding="utf-8") as f:
            json.dump(record, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ Failed to save meeting record {record.get('meeting_id')}: {e}")


def load_meeting_record(meeting_id: str) -> Optional[dict]:
    path = _record_path(meeting_id)
    if not path.exists():
        return None
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"⚠️ Failed to load meeting record {meeting_id}: {e}")
        return None


def list_meeting_records() -> list:
    if not STORE_DIR.exists():
        return []

    records = []
    for path in STORE_DIR.glob("*.json"):
        try:
            with path.open("r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"⚠️ Failed to read meeting record {path}: {e}")
            continue
        meeting_data = data.get("meeting_data") or {}
        records.append({
            "meeting_id": data.get("meeting_id"),
            "meeting_title": meeting_data.get("meeting_title", ""),
            "date": meeting_data.get("date", ""),
            "created_at": data.get("created_at", ""),
        })

    records.sort(key=lambda r: r.get("created_at") or "", reverse=True)
    return records


def update_meeting_record(meeting_id: str, meeting_data: dict) -> Optional[dict]:
    record = load_meeting_record(meeting_id)
    if record is None:
        return None
    record["meeting_data"] = meeting_data
    record["updated_at"] = datetime.now(timezone.utc).isoformat()
    save_meeting_record(record)
    return record


def delete_meeting_record(meeting_id: str) -> bool:
    path = _record_path(meeting_id)
    deleted = path.exists()
    if deleted:
        path.unlink()

    audio_dir = audio_dir_for(meeting_id)
    if audio_dir.exists():
        shutil.rmtree(audio_dir, ignore_errors=True)

    return deleted
